# -*- coding: utf-8 -*-
"""
report_builder.py
==================
Génère un rapport Word (.docx) professionnel et soigné à partir des
insights produits par smart_analysis.py : page de garde avec logo,
sommaire, résumé exécutif, sections illustrées par des graphiques
Plotly exportés en image, tableaux de données, pied de page avec
numérotation automatique.

Usage :
    from report_builder import build_word_report
    buffer = build_word_report(df, profile, insights, meta={
        "title": "Rapport d'analyse — Collecte X",
        "org_name": "CCR-B",
        "logo_path": "PP CCRB.png",
        "subtitle": "Analyse intelligente automatisée",
        "author": "CCR-B — Suivi-Évaluation",
    })
    # buffer est un io.BytesIO prêt à être proposé au téléchargement
"""

from __future__ import annotations

import io
import os
import re
from datetime import datetime
from typing import Optional

import pandas as pd

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    import plotly.io as pio
    _HAS_PLOTLY_IO = True
except Exception:  # pragma: no cover
    _HAS_PLOTLY_IO = False


GREEN_DARK = RGBColor(0x1B, 0x5E, 0x20)
GREEN_MED = RGBColor(0x2E, 0x7D, 0x32)
GREEN_LIGHT = "E8F5E9"  # hex sans # pour le shading
GREY_TEXT = RGBColor(0x37, 0x47, 0x4F)


# =============================================================================
# UTILITAIRES BAS NIVEAU (OOXML)
# =============================================================================
def _set_cell_shading(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_field(paragraph, field_code: str):
    """Insère un champ Word (ex: PAGE, NUMPAGES, TOC) qui se met à jour à l'ouverture."""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)
    return run


def add_page_number_footer(section, org_name: str):
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{org_name} — Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = GREY_TEXT
    _add_field(p, "PAGE")
    run2 = p.add_run(" / ")
    run2.font.size = Pt(9)
    run2.font.color.rgb = GREY_TEXT
    _add_field(p, "NUMPAGES")


def add_toc(doc: Document):
    """Insère un champ Sommaire (TOC) : à ouvrir dans Word, faire clic droit > Mettre à jour les champs."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Ouvrez ce document dans Word et faites clic droit → Mettre à jour les champs pour générer le sommaire."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(placeholder)
    run._r.append(fld_end)


def _add_markdown_bold_runs(paragraph, text: str, size=11, color=None):
    """Ajoute du texte à un paragraphe en interprétant les **gras** façon markdown léger."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        content = part[2:-2] if bold else part
        run = paragraph.add_run(content)
        run.bold = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color


def fig_to_png_bytes(fig, width=1000, height=560, scale=2) -> Optional[bytes]:
    if fig is None or not _HAS_PLOTLY_IO:
        return None
    try:
        return pio.to_image(fig, format="png", width=width, height=height, scale=scale)
    except Exception:
        return None


def add_dataframe_table(doc: Document, df: pd.DataFrame, max_rows: int = 15):
    df_show = df.reset_index() if df.index.name or not isinstance(df.index, pd.RangeIndex) else df
    df_show = df_show.head(max_rows)
    n_cols = len(df_show.columns)
    table = doc.add_table(rows=1, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1" if "Light Grid Accent 1" in [s.name for s in doc.styles] else "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df_show.columns):
        hdr_cells[i].text = str(col)
        _set_cell_shading(hdr_cells[i], "1B5E20")
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.bold = True
                r.font.size = Pt(9)

    for _, row in df_show.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = "" if pd.isna(val) else str(val)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table


# =============================================================================
# STYLES DU DOCUMENT
# =============================================================================
def _configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = GREY_TEXT

    for level, size, color in [(1, 20, GREEN_DARK), (2, 15, GREEN_MED), (3, 12.5, GREEN_MED)]:
        try:
            style = doc.styles[f"Heading {level}"]
            style.font.name = "Calibri"
            style.font.size = Pt(size)
            style.font.color.rgb = color
            style.font.bold = True
        except KeyError:
            pass

    section = doc.sections[0]
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)


def _add_cover_page(doc: Document, meta: dict):
    logo_path = meta.get("logo_path")
    if logo_path and os.path.exists(logo_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            run = p.add_run()
            run.add_picture(logo_path, width=Inches(1.8))
        except Exception:
            pass

    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(meta.get("title", "Rapport d'analyse de données"))
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = GREEN_DARK

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_p.add_run(meta.get("subtitle", "Analyse intelligente automatisée"))
    subtitle_run.font.size = Pt(15)
    subtitle_run.font.color.rgb = GREEN_MED
    subtitle_run.italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    org_p = doc.add_paragraph()
    org_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org_run = org_p.add_run(meta.get("org_name", ""))
    org_run.font.size = Pt(13)
    org_run.bold = True

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_p.add_run(f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}")
    date_run.font.size = Pt(11)
    date_run.font.color.rgb = GREY_TEXT

    author = meta.get("author")
    if author:
        auth_p = doc.add_paragraph()
        auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        auth_run = auth_p.add_run(author)
        auth_run.font.size = Pt(10)
        auth_run.font.color.rgb = GREY_TEXT

    doc.add_page_break()


# =============================================================================
# CONSTRUCTION DU RAPPORT
# =============================================================================
def build_word_report(df: pd.DataFrame, profile, insights: list, meta: Optional[dict] = None,
                       executive_summary: Optional[str] = None,
                       ai_interpretation: Optional[str] = None) -> io.BytesIO:
    """
    Construit le rapport Word complet et retourne un buffer BytesIO prêt
    pour st.download_button.

    Paramètres :
      df : DataFrame analysé (pour le rappel du volume de données)
      profile : DataProfile (voir smart_analysis.profile_dataframe)
      insights : liste d'Insight (voir smart_analysis.generate_smart_insights)
      meta : dict optionnel {title, subtitle, org_name, logo_path, author}
      executive_summary : texte optionnel (sinon généré automatiquement)
      ai_interpretation : texte optionnel (markdown léger) d'interprétation approfondie
                          générée par un modèle d'IA (voir ai_interpreter.py). Si fourni,
                          une section dédiée est ajoutée juste après le résumé exécutif.
    """
    meta = meta or {}
    doc = Document()
    _configure_styles(doc)
    _add_cover_page(doc, meta)

    # --- Sommaire ---
    toc_heading = doc.add_heading("Sommaire", level=1)
    add_toc(doc)
    doc.add_page_break()

    # --- Résumé exécutif ---
    doc.add_heading("Résumé exécutif", level=1)
    if executive_summary:
        for line in executive_summary.split("\n"):
            line = line.strip()
            if not line:
                continue
            p = doc.add_paragraph(style="List Bullet" if line.startswith("- ") else None)
            _add_markdown_bold_runs(p, line[2:] if line.startswith("- ") else line, size=11)
    else:
        p = doc.add_paragraph()
        _add_markdown_bold_runs(
            p,
            f"Ce rapport présente une analyse automatisée de **{profile.n_rows} enregistrements** "
            f"sur **{profile.n_cols} variables**, avec un taux de complétude de **{profile.completeness}%**.",
        )

    doc.add_paragraph()

    # --- Interprétation IA (optionnelle) ---
    if ai_interpretation:
        doc.add_heading("🤖 Interprétation approfondie (IA)", level=1)
        note_p = doc.add_paragraph()
        note_run = note_p.add_run(
            "Cette section a été générée par un modèle d'intelligence artificielle à partir des "
            "constats statistiques ci-dessus. Elle doit être relue et validée par l'équipe MEAL "
            "avant diffusion officielle."
        )
        note_run.italic = True
        note_run.font.size = Pt(9)
        note_run.font.color.rgb = GREY_TEXT
        for line in ai_interpretation.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("## ") or line.startswith("# "):
                doc.add_heading(line.lstrip("#").strip(), level=2)
            elif line.startswith(("- ", "* ")):
                p = doc.add_paragraph(style="List Bullet")
                _add_markdown_bold_runs(p, line[2:], size=11)
            elif re.match(r"^\d+\.\s", line):
                p = doc.add_paragraph(style="List Number")
                _add_markdown_bold_runs(p, re.sub(r"^\d+\.\s", "", line), size=11)
            else:
                p = doc.add_paragraph()
                _add_markdown_bold_runs(p, line, size=11)
        doc.add_paragraph()

    # --- Sections par catégorie ---
    categories_order = [
        "Qualité des données",
        "Répartitions clés",
        "Relations entre variables",
        "Tendances temporelles",
        "Répartition géographique",
    ]
    present_categories = [c for c in categories_order if any(i.category == c for i in insights)]
    # Ajouter les catégories non prévues (robustesse)
    for i in insights:
        if i.category not in present_categories:
            present_categories.append(i.category)

    for cat in present_categories:
        doc.add_heading(cat, level=1)
        cat_insights = [i for i in insights if i.category == cat]
        for ins in cat_insights:
            doc.add_heading(ins.title, level=2)
            p = doc.add_paragraph()
            _add_markdown_bold_runs(p, ins.narrative, size=11)

            img_bytes = fig_to_png_bytes(ins.fig)
            if img_bytes:
                img_stream = io.BytesIO(img_bytes)
                img_p = doc.add_paragraph()
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = img_p.add_run()
                try:
                    run.add_picture(img_stream, width=Inches(6.0))
                except Exception:
                    pass

            if ins.table is not None and isinstance(ins.table, pd.DataFrame) and len(ins.table) > 0:
                try:
                    add_dataframe_table(doc, ins.table, max_rows=12)
                except Exception:
                    pass
            doc.add_paragraph()

    # --- Pied de page avec numérotation ---
    add_page_number_footer(doc.sections[0], meta.get("org_name", "CCR-B"))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
